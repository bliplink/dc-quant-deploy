#!/usr/bin/env python3
"""Generate the final product and acceptance Word document from Markdown.

The script intentionally supports only the Markdown constructs used by the
product document.  Keeping the source of truth in Markdown makes business-rule
changes reviewable, while the generated DOCX is the user-facing deliverable.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ACCENT = "F7A600"
DARK = "171B26"
HEADER = "252B3A"
LIGHT = "F2F4F7"
MUTED = "667085"
WHITE = "FFFFFF"


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr, fld_char2))
    set_run_font(run, size=9, color=MUTED)


def add_toc(document):
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 1"]
    heading.add_run("目录")
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "在 Word 中右键此处并选择“更新域”以刷新目录。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, fallback, end))
    document.add_page_break()


def clean_inline(text):
    text = re.sub(r"!\[([^]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def add_rich_text(paragraph, text, base_size=10.5):
    # Preserve inline-code visibility without implementing a full Markdown parser.
    cursor = 0
    for match in re.finditer(r"`([^`]+)`", text):
        if match.start() > cursor:
            run = paragraph.add_run(clean_inline(text[cursor:match.start()]))
            set_run_font(run, size=base_size)
        run = paragraph.add_run(match.group(1))
        set_run_font(run, name="Consolas", size=max(8.5, base_size - 1), color=DARK)
        run.font.highlight_color = None
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(clean_inline(text[cursor:]))
        set_run_font(run, size=base_size)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(5)

    heading_specs = {
        "Heading 1": (17, DARK, Pt(14), Pt(7)),
        "Heading 2": (14, DARK, Pt(11), Pt(5)),
        "Heading 3": (11.5, ACCENT, Pt(8), Pt(4)),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Cm(0.45)
    code_style.paragraph_format.right_indent = Cm(0.25)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DC SaaS 加密货币永续合约系统｜产品与验收说明")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def metadata_value(markdown, *labels, default="-"):
    for label in labels:
        match = re.search(rf"^{re.escape(label)}：(.+)$", markdown, re.MULTILINE)
        if match:
            return clean_inline(match.group(1)).strip()
    return default


def add_cover(document, markdown):
    markdown_title = next(
        (line[2:].strip() for line in markdown.splitlines() if line.startswith("# ")),
        "DC SaaS 加密货币永续合约系统",
    )
    if "多租户控制面" in markdown_title:
        cover_title = "DC SaaS 加密货币\n多租户控制面"
        cover_subtitle = "业务能力、租户边界、数据流与生产验收报告"
    else:
        cover_title = "DC SaaS 加密货币\n永续合约交易系统"
        cover_subtitle = "产品说明、业务规则、数据流与完整验收报告"
    version = metadata_value(markdown, "文档版本", default="V1.0")
    baseline = metadata_value(markdown, "基线日期", "验收日期", default="2026-08-30")
    branch = metadata_value(markdown, "代码分支", default="saas-crypto")
    environment = metadata_value(
        markdown,
        "当前生产验证环境",
        "生产验证环境",
        default="18.140.45.126 / dc-saas",
    )
    environment_ip = re.search(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", environment)
    compose_project = re.search(r"Compose 项目\s+([\w-]+)", environment)
    if environment_ip:
        environment = environment_ip.group(0)
        if compose_project:
            environment += f" / {compose_project.group(1)}"

    for _ in range(4):
        document.add_paragraph()
    tag = document.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("PRODUCT · BUSINESS RULES · DATA FLOW · ACCEPTANCE")
    set_run_font(run, name="Arial", size=10, bold=True, color=ACCENT)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run(cover_title)
    set_run_font(run, size=28, bold=True, color=DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(cover_subtitle)
    set_run_font(run, size=14, color=MUTED)

    document.add_paragraph()
    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, name="Arial", size=12, color=ACCENT)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(26)
    run = meta.add_run(
        f"版本：{version}\n"
        f"基线：{baseline}\n"
        f"分支：{branch}\n"
        f"验收环境：{environment}"
    )
    set_run_font(run, size=11, color=DARK)
    document.add_page_break()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        cells = [clean_inline(c) for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(document, rows):
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for r_index, row in enumerate(rows):
        for c_index in range(width):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = row[c_index] if c_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(
                run,
                size=8.5,
                bold=(r_index == 0),
                color=WHITE if r_index == 0 else DARK,
            )
            if r_index == 0:
                set_cell_shading(cell, HEADER)
            elif r_index % 2 == 0:
                set_cell_shading(cell, LIGHT)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(document, markdown_path, alt, source_dir):
    image_path = (source_dir / markdown_path).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"[图片不存在：{image_path}]")
        set_run_font(run, color="C00000")
        return
    with Image.open(image_path) as image:
        ratio = image.width / image.height
    max_width = 6.45
    max_height = 8.2
    width = min(max_width, max_height * ratio)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(f"图：{alt}")
    set_run_font(run, size=8.5, color=MUTED)


def add_markdown(document, markdown, source_dir):
    lines = markdown.splitlines()
    # The Markdown title and four metadata lines are represented on the cover.
    body_start = next((i for i, line in enumerate(lines) if line.startswith("## ")), 0)
    lines = lines[body_start:]
    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                paragraph = document.add_paragraph(style="Code Block")
                paragraph.paragraph_format.keep_together = True
                run = paragraph.add_run("\n".join(code_lines))
                set_run_font(run, name="Consolas", size=8.5, color=DARK)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(document, rows)
            continue
        image_match = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(document, image_match.group(2), image_match.group(1), source_dir)
            i += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = min(3, len(heading.group(1)) - 1)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_rich_text(paragraph, heading.group(2), 17 if level == 1 else 14 if level == 2 else 11.5)
            i += 1
            continue
        ordered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            add_rich_text(paragraph, ordered.group(2))
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_rich_text(paragraph, bullet.group(1))
            i += 1
            continue
        paragraph = document.add_paragraph()
        add_rich_text(paragraph, stripped)
        i += 1


def generate(source, output):
    markdown = source.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    add_cover(document, markdown)
    add_toc(document)
    add_markdown(document, markdown, source.parent)
    document.core_properties.title = "DC SaaS 加密货币永续合约系统产品与验收说明"
    document.core_properties.subject = "业务规则、数据流、测试与截图证据"
    document.core_properties.author = "DC SaaS Crypto 项目组"
    document.core_properties.keywords = "SaaS, 永续合约, 撮合, 强平, ADL, 多租户"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=Path(__file__).resolve().parents[1]
        / "docs"
        / "DC_SAAS_CRYPTO_PRODUCT_AND_ACCEPTANCE_20260829.zh-CN.md",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).resolve().parents[1]
        / "docs"
        / "DC_SAAS_CRYPTO_PRODUCT_AND_ACCEPTANCE_20260829.zh-CN.docx",
    )
    args = parser.parse_args()
    generate(args.source.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
